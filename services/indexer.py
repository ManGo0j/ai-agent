import os
import glob
import uuid
import hashlib
import pdfplumber
import asyncio
from docx import Document
from datetime import datetime
from typing import List, Generator

from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http import models
from fastembed import SparseTextEmbedding

# Исправленные импорты для SQLAlchemy 2.0
from sqlalchemy import select, text
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker

# Импорт моделей БД
from database.models import Base, AdminDocument, DocumentChunk

# Импорт инструментов для обработки текста
from langchain_text_splitters import RecursiveCharacterTextSplitter

load_dotenv()

# --- КОНФИГУРАЦИЯ ---
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(BASE_DIR, "docs")
# Указываем путь к кэшу моделей внутри контейнера
MODEL_CACHE_DIR = os.path.join(BASE_DIR, "model_cache") 

# Настройки Qdrant и БД
QDRANT_URL = os.getenv("QDRANT_URL", "http://accountant_qdrant:6333")
DATABASE_URL = os.getenv("DATABASE_URL")
COLLECTION_NAME = "knowledge_base"
MODEL_NAME = "all-MiniLM-L6-v2"

# Настройки разбиения текста
CHUNK_SIZE = 2000
CHUNK_OVERLAP = 500

# --- ИНИЦИАЛИЗАЦИЯ ---
qdrant_client = QdrantClient(url=QDRANT_URL)

# Модель для плотных векторов
model = SentenceTransformer(MODEL_NAME)

# ИСПРАВЛЕНИЕ: Добавлен cache_dir для предотвращения ошибки NoSuchFile
sparse_model = SparseTextEmbedding(
    model_name="prithivida/Splade_PP_en_v1",
    cache_dir=MODEL_CACHE_DIR
)

engine = create_async_engine(DATABASE_URL)

# Создание фабрики асинхронных сессий
AsyncSessionLocal = async_sessionmaker(
    bind=engine, 
    class_=AsyncSession, 
    expire_on_commit=False
)

# Инициализируем умный сплиттер
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=["\n\n", "\n", " ", ""]
)

# --- ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ---

def get_file_hash(file_path: str) -> str:
    """Вычисляет хеш файла для проверки на дубликаты."""
    hasher = hashlib.sha256()
    with open(file_path, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hasher.update(chunk)
    return hasher.hexdigest()

def stream_extract_docx(file_path: str, batch_paragraphs: int = 30) -> Generator[str, None, None]:
    """Потоковое чтение DOCX по порциям абзацев."""
    try:
        doc = Document(file_path)
        buffer = []
        for para in doc.paragraphs:
            if para.text.strip():
                buffer.append(para.text)
            if len(buffer) >= batch_paragraphs:
                yield "\n".join(buffer)
                buffer = []
        if buffer:
            yield "\n".join(buffer)
    except Exception as e:
        print(f"Ошибка чтения DOCX {file_path}: {e}")

def stream_extract_pdf(file_path: str) -> Generator[str, None, None]:
    """Потоковое чтение PDF по страницам."""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    yield text
    except Exception as e:
        print(f"Ошибка чтения PDF {file_path}: {e}")

async def process_files():
    """Основная функция обработки документов с поддержкой потоков."""
    
    # 1. Проверка или создание коллекции
    try:
        if not qdrant_client.collection_exists(COLLECTION_NAME):
            qdrant_client.create_collection(
                collection_name=COLLECTION_NAME,
                vectors_config={
                    "default": models.VectorParams(
                        size=384, 
                        distance=models.Distance.COSINE
                    )
                },
                sparse_vectors_config={
                    "sparse-text": models.SparseVectorParams(
                        index=models.SparseIndexParams(
                            on_disk=False,
                        )
                    )
                }
            )
            print(f"Создана гибридная коллекция: {COLLECTION_NAME}")
        else:
            print(f"Коллекция {COLLECTION_NAME} готова к работе")
    except Exception as e:
        print(f"Ошибка Qdrant: {e}")
        return

    # 2. Поиск файлов
    files = []
    for ext in ['*.pdf', '*.docx', '*.doc', '*.txt']:
        files.extend(glob.glob(os.path.join(DOCS_DIR, ext)))

    if not files:
        print("Документы для обработки не найдены")
        return

    print(f"Найдено файлов: {len(files)}")

    # 3. Основной цикл
    async with AsyncSessionLocal() as session:
        for file_path in files:
            file_name = os.path.basename(file_path)
            
            try:
                file_hash = get_file_hash(file_path)
                result = await session.execute(
                    select(AdminDocument).where(AdminDocument.file_hash == file_hash)
                )
                if result.scalar_one_or_none():
                    print(f"Файл {file_name} уже есть в базе")
                    continue

                print(f"Обработка: {file_name}")

                doc_type = "word" if file_name.endswith(('.docx', '.doc')) else "pdf"
                if file_name.endswith('.txt'): doc_type = "text"

                db_doc = AdminDocument(
                    document_name=file_name,
                    file_path=file_path,
                    file_hash=file_hash,
                    document_type=doc_type,
                    upload_date=datetime.utcnow()
                )
                session.add(db_doc)
                await session.flush()
                
                # Сохраняем ID, чтобы не потерять его после expunge_all
                current_doc_id = db_doc.id

                # Выбор метода чтения
                if doc_type == "pdf":
                    content_generator = stream_extract_pdf(file_path)
                elif doc_type == "word":
                    content_generator = stream_extract_docx(file_path)
                else:
                    def text_gen():
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            yield f.read()
                    content_generator = text_gen()

                total_chunks = 0
                for text_segment in content_generator:
                    if not text_segment.strip():
                        continue

                    # Нарезка сегмента на чанки
                    segment_chunks = text_splitter.split_text(text_segment)
                    
                    contextual_chunks = []
                    for chunk in segment_chunks:
                        header = f"ИСТОЧНИК: {file_name}\nТИП: {doc_type}\n\n"
                        contextual_chunks.append(header + chunk)

                    if not contextual_chunks:
                        continue

                    # ОПТИМИЗАЦИЯ: Вынос тяжелых вычислений в отдельные потоки
                    dense_vecs = await asyncio.to_thread(model.encode, contextual_chunks)
                    sparse_vecs = await asyncio.to_thread(lambda: list(sparse_model.embed(contextual_chunks)))

                    points = []
                    for i, (txt, d_vec, s_vec) in enumerate(zip(contextual_chunks, dense_vecs, sparse_vecs)):
                        point_id = str(uuid.uuid4())
                        
                        db_chunk = DocumentChunk(
                            document_id=current_doc_id, # Используем сохраненный ID
                            chunk_index=total_chunks + i,
                            chunk_text=txt,
                            embedding_id=point_id
                        )
                        session.add(db_chunk)

                        points.append(models.PointStruct(
                            id=point_id,
                            vector={
                                "default": d_vec.tolist(),
                                "sparse-text": models.SparseVector(
                                    indices=s_vec.indices.tolist(),
                                    values=s_vec.values.tolist()
                                )
                            },
                            payload={
                                "document_id": current_doc_id,
                                "document_name": file_name,
                                "text": txt
                            }
                        ))

                    # Загрузка в Qdrant
                    if points:
                        qdrant_client.upsert(
                            collection_name=COLLECTION_NAME,
                            points=points
                        )
                    
                    total_chunks += len(contextual_chunks)
                    
                    # Фиксация и очистка сессии для экономии памяти
                    await session.commit()
                    session.expunge_all()

                print(f"Готово: {file_name} (чанков: {total_chunks})")

            except Exception as e:
                await session.rollback()
                print(f"Ошибка в файле {file_name}: {e}")
                continue

if __name__ == "__main__":
    asyncio.run(process_files())